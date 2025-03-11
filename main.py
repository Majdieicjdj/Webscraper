# from flask import Flask, request, jsonify, send_file
# from flask_cors import CORS
# import os
# import time
# import logging
# import subprocess
# import json

# app = Flask(__name__)
# CORS(app)  # Enable CORS

# logging.basicConfig(level=logging.INFO)
# os.makedirs('downloads', exist_ok=True)

# scrape_sessions = {}

# # Set Chromium executable path
# CHROMIUM_PATH = "C:\\chromium\\chrome-win\\chrome.exe"

# # Ensure Chromium exists
# if not os.path.exists(CHROMIUM_PATH):
#     raise FileNotFoundError(f"Chromium executable not found at {CHROMIUM_PATH}. Please check the path.")

# def scrape_website_subprocess(url, elements=None):
#     """Runs Pyppeteer in a separate subprocess to avoid threading conflicts."""
#     try:
#         script = f"""
# import asyncio
# from pyppeteer import launch
# from bs4 import BeautifulSoup
# import time
# import json

# async def scrape():
#     browser = await launch(
#         headless=True,
#         executablePath=r"{CHROMIUM_PATH}",
#         args=['--no-sandbox', '--disable-setuid-sandbox']
#     )
#     page = await browser.newPage()
#     await page.goto("{url}", {{'waitUntil': 'networkidle2', 'timeout': 30000}})
#     await asyncio.sleep(2)

#     content = await page.content()
#     soup = BeautifulSoup(content, 'html.parser')

#     results = {{
#         'url': "{url}",
#         'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
#         'title': await page.title(),
#         'data': {{}}
#     }}

#     elements = {json.dumps(elements)}

#     if elements:
#         for element in elements:
#             try:
#                 elements_data = await page.evaluate('(selector) => Array.from(document.querySelectorAll(selector)).map(el => el.innerText)', element)
#                 results['data'][element] = elements_data
#             except Exception:
#                 results['data'][element] = "Error extracting element"
#     else:
#         results['data']['text'] = soup.get_text(separator='\\n', strip=True)

#     await browser.close()
#     print(json.dumps(results))  # Print JSON output

# asyncio.run(scrape())
# """
#         process = subprocess.run(
#             ["python", "-c", script],
#             capture_output=True,
#             text=True
#         )
#         output = process.stdout.strip()
#         return json.loads(output) if output else {"error": "Failed to scrape"}
#     except Exception as e:
#         logging.error(f"Scraping error: {str(e)}")
#         return {'error': str(e)}

# @app.route('/api/scrape', methods=['POST'])
# def scrape():
#     data = request.json
#     url = data.get('url')
#     elements = data.get('elements', [])

#     if not url:
#         return jsonify({'error': 'URL is required'}), 400

#     try:
#         results = scrape_website_subprocess(url, elements)

#         session_id = int(time.time())
#         scrape_sessions[session_id] = results

#         return jsonify({'success': True, 'data': results, 'session_id': session_id})
    
#     except Exception as e:
#         logging.error(f"Scraping failed: {str(e)}")
#         return jsonify({'error': f'Scraping failed: {str(e)}'}), 500

# @app.route('/api/download/<format>/<int:session_id>')
# def download(format, session_id):
#     if session_id not in scrape_sessions:
#         return jsonify({'error': 'Session not found'}), 404

#     filename = f"scraped_data.{format}"
#     return send_file(filename, as_attachment=True)

# if __name__ == "__main__":
#     app.run(host="0.0.0.0", port=5000, debug=False, use_reloader=False)  # Ensures stability





# from flask import Flask, request, jsonify, send_file
# from flask_cors import CORS
# import os
# import time
# import logging
# import subprocess
# import json
# import docx
# from fpdf import FPDF

# app = Flask(__name__)
# CORS(app)  # Enable CORS

# logging.basicConfig(level=logging.INFO)
# os.makedirs('downloads', exist_ok=True)

# scrape_sessions = {}

# # Set Chromium executable path
# CHROMIUM_PATH = "C:\\chromium\\chrome-win\\chrome.exe"

# # Ensure Chromium exists
# if not os.path.exists(CHROMIUM_PATH):
#     raise FileNotFoundError(f"Chromium executable not found at {CHROMIUM_PATH}. Please check the path.")

# def scrape_website_subprocess(url, elements=None):
#     """Runs Pyppeteer in a separate subprocess to avoid threading conflicts."""
#     try:
#         script = f"""
# import asyncio
# from pyppeteer import launch
# from bs4 import BeautifulSoup
# import time
# import json

# async def scrape():
#     browser = await launch(
#         headless=True,
#         executablePath=r"{CHROMIUM_PATH}",
#         args=['--no-sandbox', '--disable-setuid-sandbox']
#     )
#     page = await browser.newPage()
#     await page.goto("{url}", {{'waitUntil': 'networkidle2', 'timeout': 30000}})
#     await asyncio.sleep(2)

#     content = await page.content()
#     soup = BeautifulSoup(content, 'html.parser')

#     results = {{
#         'url': "{url}",
#         'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
#         'title': await page.title(),
#         'data': {{}}
#     }}

#     elements = {json.dumps(elements)}

#     if elements:
#         for element in elements:
#             try:
#                 elements_data = await page.evaluate('(selector) => Array.from(document.querySelectorAll(selector)).map(el => el.innerText)', element)
#                 results['data'][element] = elements_data
#             except Exception:
#                 results['data'][element] = "Error extracting element"
#     else:
#         results['data']['text'] = soup.get_text(separator='\\n', strip=True)

#     await browser.close()
#     print(json.dumps(results))  # Print JSON output

# asyncio.run(scrape())
# """
#         process = subprocess.run(
#             ["python", "-c", script],
#             capture_output=True,
#             text=True
#         )
#         output = process.stdout.strip()
#         return json.loads(output) if output else {"error": "Failed to scrape"}
#     except Exception as e:
#         logging.error(f"Scraping error: {str(e)}")
#         return {'error': str(e)}

# @app.route('/api/scrape', methods=['POST'])
# def scrape():
#     data = request.json
#     url = data.get('url')
#     elements = data.get('elements', [])

#     if not url:
#         return jsonify({'error': 'URL is required'}), 400

#     try:
#         results = scrape_website_subprocess(url, elements)

#         session_id = int(time.time())
#         scrape_sessions[session_id] = results

#         return jsonify({'success': True, 'data': results, 'session_id': session_id})
    
#     except Exception as e:
#         logging.error(f"Scraping failed: {str(e)}")
#         return jsonify({'error': f'Scraping failed: {str(e)}'}), 500

# def generate_word_file(session_id):
#     """Generates a Word (.docx) file from the scraped data."""
#     if session_id not in scrape_sessions:
#         return None

#     data = scrape_sessions[session_id]
#     filename = f"downloads/scraped_data_{session_id}.docx"

#     doc = docx.Document()
#     doc.add_heading('Scraped Data', level=1)
#     doc.add_paragraph(f"URL: {data.get('url', 'N/A')}")
#     doc.add_paragraph(f"Time: {data.get('timestamp', 'N/A')}")
#     doc.add_paragraph("Extracted Data:\n")

#     for key, value in data.get('data', {}).items():
#         doc.add_paragraph(f"{key}: {value}\n")

#     doc.save(filename)
#     return filename

# def generate_pdf_file(session_id):
#     """Generates a PDF file from the scraped data."""
#     if session_id not in scrape_sessions:
#         return None

#     data = scrape_sessions[session_id]
#     filename = f"downloads/scraped_data_{session_id}.pdf"

#     pdf = FPDF()
#     pdf.set_auto_page_break(auto=True, margin=15)
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     pdf.cell(200, 10, txt="Scraped Data", ln=True, align='C')
#     pdf.ln(10)
#     pdf.cell(200, 10, txt=f"URL: {data.get('url', 'N/A')}", ln=True)
#     pdf.cell(200, 10, txt=f"Time: {data.get('timestamp', 'N/A')}", ln=True)
#     pdf.ln(10)
    
#     for key, value in data.get('data', {}).items():
#         pdf.multi_cell(0, 10, f"{key}: {value}\n")

#     pdf.output(filename)
#     return filename

# @app.route('/api/download/<format>/<int:session_id>')
# def download(format, session_id):
#     """Handles file downloads for Word and PDF formats."""
#     if session_id not in scrape_sessions:
#         return jsonify({'error': 'Session not found'}), 404

#     if format == "docx":
#         filename = generate_word_file(session_id)
#     elif format == "pdf":
#         filename = generate_pdf_file(session_id)
#     else:
#         return jsonify({'error': 'Invalid format'}), 400

#     if not filename or not os.path.exists(filename):
#         return jsonify({'error': 'File generation failed'}), 500

#     return send_file(filename, as_attachment=True)

# if __name__ == "__main__":
#     app.run(host="0.0.0.0", port=5000, debug=False, use_reloader=False)  # Ensures stability







from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import time
import logging
import subprocess
import json
import docx
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches
from fpdf import FPDF
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin

app = Flask(__name__)
CORS(app)  # Enable CORS

logging.basicConfig(level=logging.INFO)
os.makedirs('downloads', exist_ok=True)

scrape_sessions = {}

CHROMIUM_PATH = "C:\\chromium\\chrome-win\\chrome.exe"

if not os.path.exists(CHROMIUM_PATH):
    logging.warning(f"Chromium executable not found at {CHROMIUM_PATH}. Path will be used anyway.")

def scrape_website(url, elements=None):
    """Extracts text, images, and links properly formatted."""
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        response = requests.get(url, headers=headers)
        response.raise_for_status()

        soup = BeautifulSoup(response.text, "html.parser")
        results = {
            'url': url,
            'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
            'title': soup.title.string if soup.title else "No Title",
            'data': {'text': soup.get_text(separator="\n", strip=True)},
            'images': [],
            'links': []
        }

        # Extract images
        for img in soup.find_all("img"):
            img_url = urljoin(url, img.get("src", ""))
            results["images"].append(img_url)

        # Extract links with meaningful text
        for link in soup.find_all("a", href=True):
            link_text = link.get_text(strip=True) or link["href"]
            results["links"].append({"text": link_text, "url": urljoin(url, link["href"])})

        return results
    except Exception as e:
        logging.error(f"Scraping error: {str(e)}")
        return {'error': str(e)}

@app.route('/api/scrape', methods=['POST'])
def scrape():
    data = request.json
    url = data.get('url')

    if not url:
        return jsonify({'error': 'URL is required'}), 400

    results = scrape_website(url)
    session_id = int(time.time())
    scrape_sessions[session_id] = results

    return jsonify({'success': True, 'data': results, 'session_id': session_id})

def generate_word_file(session_id):
    """Generates a Word (.docx) file with proper formatting and embedded images."""
    try:
        if session_id not in scrape_sessions:
            return None

        data = scrape_sessions[session_id]
        filename = f"downloads/scraped_data_{session_id}.docx"

        doc = docx.Document()
        doc.add_heading('Scraped Data Report', level=1).alignment = 1

        doc.add_paragraph(f"URL: {data.get('url', 'N/A')}")
        doc.add_paragraph(f"Page Title: {data.get('title', 'N/A')}")
        doc.add_paragraph(f"Timestamp: {data.get('timestamp', 'N/A')}")
        doc.add_paragraph("=" * 50)

        if 'error' in data:
            error_para = doc.add_paragraph()
            run = error_para.add_run("ERROR: ")
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
            error_para.add_run(data['error'])
        else:
            doc.add_heading('Extracted Text', level=2)
            doc.add_paragraph(data['data']['text'])

            # Add links
            doc.add_heading("Extracted Links", level=2)
            for link in data.get('links', []):
                p = doc.add_paragraph()
                p.add_run(link["text"]).bold = True
                p.add_run(f" ({link['url']})").italic = True

            # Add images
            doc.add_heading("Extracted Images", level=2)
            for img_url in data.get('images', []):
                try:
                    img_data = requests.get(img_url, timeout=5).content
                    img_filename = f"downloads/temp_image_{session_id}.jpg"
                    with open(img_filename, "wb") as img_file:
                        img_file.write(img_data)
                    doc.add_picture(img_filename, width=Inches(4))
                except Exception as img_err:
                    doc.add_paragraph(f"Could not load image: {img_url}")

        doc.save(filename)
        return filename
    except Exception as e:
        logging.error(f"Word file generation error: {str(e)}")
        return None

def generate_pdf_file(session_id):
    """Generates a PDF file with proper formatting."""
    try:
        if session_id not in scrape_sessions:
            return None

        data = scrape_sessions[session_id]
        filename = f"downloads/scraped_data_{session_id}.pdf"

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", 'B', size=16)
        pdf.cell(200, 10, txt="Scraped Data Report", ln=True, align='C')
        pdf.ln(5)

        pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 10, txt=f"URL: {data.get('url', 'N/A')}")
        pdf.multi_cell(0, 10, txt=f"Page Title: {data.get('title', 'N/A')}")
        pdf.multi_cell(0, 10, txt=f"Timestamp: {data.get('timestamp', 'N/A')}")

        pdf.ln(5)
        pdf.cell(200, 0, txt="=" * 80, ln=True)
        pdf.ln(10)

        if 'error' in data:
            pdf.set_text_color(255, 0, 0)
            pdf.set_font("Arial", 'B', size=14)
            pdf.cell(200, 10, txt="ERROR", ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", size=11)
            pdf.multi_cell(0, 10, txt=data['error'])
        else:
            pdf.set_font("Arial", 'B', size=14)
            pdf.cell(200, 10, txt="Extracted Text", ln=True)
            pdf.ln(5)
            pdf.set_font("Arial", size=11)
            pdf.multi_cell(0, 10, txt=data['data']['text'])

            pdf.ln(5)
            pdf.set_font("Arial", 'B', size=14)
            pdf.cell(200, 10, txt="Extracted Links", ln=True)
            pdf.set_font("Arial", size=11)
            for link in data.get('links', []):
                pdf.multi_cell(0, 10, txt=f"{link['text']} ({link['url']})")

        pdf.output(filename)
        return filename
    except Exception as e:
        logging.error(f"PDF file generation error: {str(e)}")
        return None

@app.route('/api/download/<format>/<int:session_id>')
def download(format, session_id):
    """Handles file downloads."""
    filename = generate_word_file(session_id) if format == "docx" else generate_pdf_file(session_id)
    if not filename or not os.path.exists(filename):
        return jsonify({'error': 'File generation failed'}), 500
    return send_file(filename, as_attachment=True)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=False, use_reloader=False)
